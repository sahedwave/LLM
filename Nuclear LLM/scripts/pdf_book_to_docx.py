import os
import re
import subprocess
import sys
import tempfile
from multiprocessing import Pool, cpu_count
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / ".codex_tmp"))

from docx import Document


PDF_PATH = Path("/Users/VIP/Downloads/introduction-to-nuclear-engineering_compress.pdf")
OUTPUT_PATH = Path("/Users/VIP/Documents/New project/Nuclear LLM/introduction-to-nuclear-engineering_lines_no_equations_figures.docx")
PDFTOPPM = "/opt/homebrew/bin/pdftoppm"
TESSERACT = "/opt/homebrew/bin/tesseract"
TOTAL_PAGE_COUNT = 702
WORKERS = max(1, min(4, cpu_count() or 1))

HEADER_RE = re.compile(r"^\s*[\dS$§\.]+(?:\s+[A-Za-z][A-Za-z&,-]*)+\s+\d+\s*$")
PAGE_ONLY_RE = re.compile(r"^\s*\d+\s*$")
CAPTION_RE = re.compile(r"^\s*(fig(?:ure)?\.?|table)\b", re.IGNORECASE)
EQUATION_NUM_RE = re.compile(r"\(\d+\.\d+\)")
COMMON_WORD_RE = re.compile(
    r"\b(the|and|for|with|from|that|this|into|when|where|which|their|therefore|because|while|other|these|those|have|been|also|then|than|about|under|within|energy|nuclear|reactor|neutron|particle|water|matter)\b",
    re.IGNORECASE,
)


def render_page(page_number: int, temp_dir: Path) -> Path:
    stem = temp_dir / f"page_{page_number:04d}"
    subprocess.run(
        [
            PDFTOPPM,
            "-f",
            str(page_number),
            "-l",
            str(page_number),
            "-r",
            "180",
            "-singlefile",
            "-png",
            str(PDF_PATH),
            str(stem),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return stem.with_suffix(".png")


def ocr_lines(image_path: Path) -> list[str]:
    result = subprocess.run(
        [
            TESSERACT,
            str(image_path),
            "stdout",
            "-l",
            "eng",
            "--psm",
            "3",
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    return [clean_line(line) for line in result.stdout.splitlines() if clean_line(line)]


def clean_line(text: str) -> str:
    text = text.replace("|", "I")
    text = text.replace("ﬁ", "fi").replace("ﬂ", "fl")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def is_body_like(text: str) -> bool:
    words = re.findall(r"[A-Za-z]{2,}", text)
    return len(words) >= 7 and any(word[0].islower() for word in re.findall(r"[A-Za-z]+", text))


def symbol_ratio(text: str) -> float:
    if not text:
        return 0.0
    symbols = sum(1 for char in text if not char.isalnum() and not char.isspace())
    return symbols / len(text)


def digit_ratio(text: str) -> float:
    if not text:
        return 0.0
    digits = sum(1 for char in text if char.isdigit())
    return digits / len(text)


def alpha_ratio(text: str) -> float:
    if not text:
        return 0.0
    alpha = sum(1 for char in text if char.isalpha())
    return alpha / len(text)


def has_many_single_tokens(text: str) -> bool:
    tokens = text.split()
    if not tokens:
        return False
    single = sum(1 for token in tokens if len(re.sub(r"[^A-Za-z0-9]", "", token)) <= 1)
    return single >= max(3, len(tokens) // 2)


def is_equation_like(text: str) -> bool:
    words = text.split()
    lower = text.lower()
    if "=" in text and len(words) <= 18:
        return True
    if text.startswith(("x(", "E=", "S=", "Fission rate =")):
        return True
    if EQUATION_NUM_RE.search(text) and (symbol_ratio(text) > 0.09 or has_many_single_tokens(text)):
        return True
    if len(words) <= 10 and digit_ratio(text) > 0.12 and symbol_ratio(text) > 0.06 and not COMMON_WORD_RE.search(lower):
        return True
    if len(words) <= 8 and has_many_single_tokens(text) and not COMMON_WORD_RE.search(lower):
        return True
    if alpha_ratio(text) < 0.55 and symbol_ratio(text) > 0.08:
        return True
    if len(words) <= 12 and any(token in text for token in [" dE", "MeV", "keV", "MW", "joule", "fissions/day"]) and digit_ratio(text) > 0.08:
        return True
    return False


def is_figure_or_table_like(text: str) -> bool:
    stripped = text.strip()
    lower = stripped.lower()
    if not stripped:
        return True
    if PAGE_ONLY_RE.match(stripped):
        return True
    if HEADER_RE.match(stripped):
        return True
    if CAPTION_RE.match(stripped):
        return True
    if stripped.lower().startswith(("chapter ", "page ")):
        return True
    if "fig." in lower and digit_ratio(stripped) > 0.15:
        return True
    if len(stripped.split()) <= 3 and alpha_ratio(stripped) < 0.65:
        return True
    if not any(char.islower() for char in stripped) and len(stripped.split()) <= 8:
        return True
    if len(stripped.split()) <= 6 and digit_ratio(stripped) > 0.18:
        return True
    if len(stripped.split()) <= 8 and lower.count("curve") >= 1:
        return True
    if any(token in lower for token in ["ion paits", "distance from end", "mean range", "curve 4", "curve 8", "curve c"]):
        return True
    if len(stripped.split()) <= 4 and any(token in lower for token in ["mev", "kev", "mw", "barns"]):
        return True
    return False


def filter_page(lines: list[str]) -> list[str]:
    caption_or_figure_page = any(CAPTION_RE.match(line) for line in lines) and sum(
        1 for line in lines if is_body_like(line)
    ) <= max(5, len(lines) // 3)

    kept: list[str] = []
    skipping_table_or_figure_block = False
    for text in lines:
        if not text:
            continue
        lower = text.lower()
        if CAPTION_RE.match(text) or ("fig." in lower and digit_ratio(text) > 0.15):
            skipping_table_or_figure_block = True
            continue
        if skipping_table_or_figure_block:
            if is_body_like(text):
                skipping_table_or_figure_block = False
            else:
                continue
        if is_figure_or_table_like(text):
            continue
        if is_equation_like(text):
            continue
        if caption_or_figure_page and not is_body_like(text):
            continue
        kept.append(text.lstrip(". ").strip())
    return kept


def process_page(page_number: int) -> tuple[int, list[str]]:
    with tempfile.TemporaryDirectory(prefix=f"ocr_page_{page_number:04d}_") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        image_path = render_page(page_number, temp_dir)
        lines = ocr_lines(image_path)
        return page_number, filter_page(lines)


def build_document(pages: list[tuple[int, list[str]]], output_path: Path) -> None:
    document = Document()
    document.add_heading("Introduction to Nuclear Engineering", level=1)
    document.add_paragraph("OCR export with line breaks preserved where possible. Equation-like lines and figure/table text were filtered heuristically.")

    for page_number, lines in pages:
        document.add_heading(f"Page {page_number}", level=2)
        if not lines:
            document.add_paragraph("[No retained text after filtering]")
            continue
        for line in lines:
            document.add_paragraph(line)

    document.save(str(output_path))


def main() -> None:
    start_page = int(os.environ.get("START_PAGE", "1"))
    end_page = int(os.environ.get("END_PAGE", str(TOTAL_PAGE_COUNT)))
    output_path = Path(os.environ.get("OUTPUT_PATH", str(OUTPUT_PATH)))
    page_numbers = list(range(start_page, end_page + 1))
    pages: dict[int, list[str]] = {}

    with Pool(processes=WORKERS) as pool:
        for page_number, lines in pool.imap_unordered(process_page, page_numbers):
            pages[page_number] = lines
            print(f"processed page {page_number}/{end_page}", flush=True)

    ordered_pages = [(page_number, pages.get(page_number, [])) for page_number in page_numbers]
    build_document(ordered_pages, output_path)
    print(f"saved {output_path}")


if __name__ == "__main__":
    main()
